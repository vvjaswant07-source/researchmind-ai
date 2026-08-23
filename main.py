import os
import uuid
import tempfile
import json
import re
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()

# ============================================================
# ENVIRONMENT
# ============================================================

HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

print("Hugging Face key loaded:", bool(HUGGINGFACE_API_KEY))
print("Groq key loaded:", bool(GROQ_API_KEY))
print("Tavily key loaded:", bool(TAVILY_API_KEY))

if not GROQ_API_KEY:
    raise RuntimeError("GROQ_API_KEY is missing from the .env file")


# ============================================================
# IMPORTS
# ============================================================

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

from groq import Groq

try:
    from tavily import TavilyClient
except ImportError:
    TavilyClient = None

from pypdf import PdfReader

try:
    import fitz
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt


# ============================================================
# CLIENTS
# ============================================================

groq_client = Groq(api_key=GROQ_API_KEY)

if TAVILY_API_KEY and TavilyClient:
    tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
else:
    tavily_client = None
    print("WARNING: Tavily web search is disabled.")


# ============================================================
# FASTAPI
# ============================================================

app = FastAPI(
    title="ResearchMind AI",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# DATA MODELS
# ============================================================

class AgentRequest(BaseModel):
    agent: str
    prompt: str
    document_id: str | None = None


class KnowledgeRequest(BaseModel):
    title: str
    content: str


class ReportExportRequest(BaseModel):
    content: str
    title: str = "ResearchMind AI Report"
    format: str = "pdf"


class MultiAgentRequest(BaseModel):
    agents: list[str]
    prompt: str
    document_id: str | None = None


# ============================================================
# STORAGE
# ============================================================

pdf_store = {}

knowledge_store = []


# ============================================================
# AGENT INSTRUCTIONS
# ============================================================

AGENT_INSTRUCTIONS = {

    "Research Agent": """
You are the Research Agent.

Your job is web research and evidence collection.

Use web research when current information is needed.

Compare multiple sources.

Identify important findings.

Clearly separate evidence from uncertainty.

Always include source URLs supplied by the search tool.
""",

    "PDF Analysis Agent": """
You are the PDF Analysis Agent.

Analyze the uploaded PDF content supplied to you.

Answer questions using the extracted PDF content.

Do not invent information that is not present in the PDF.

If the PDF does not contain enough information, clearly say so.
""",

    "Knowledge Agent": """
You are the Knowledge Agent.

Use the private ResearchMind knowledge base when relevant.

Do not invent knowledge-base content.

Clearly distinguish knowledge-base information from general knowledge when appropriate.
""",

    "Report Writer Agent": """
You are the Report Writer Agent.

Create professional research reports.

Preferred structure:

Title

Executive Summary

1. Introduction
2. Background
3. Objectives
4. Methodology
5. Detailed Analysis
6. Findings
7. Discussion
8. Advantages / Applications
9. Limitations
10. Future Scope
11. Conclusion
12. References

Use clear Markdown.

Do not invent facts or references.
""",

    "Citation Agent": """
You are the Citation Agent.

Create accurate references only from sources actually supplied.

Never invent authors, titles, dates, publishers, or URLs.

Use APA-style references when appropriate.
""",

    "Fact Checker Agent": """
You are the Fact Checker Agent.

Check claims against supplied evidence.

For important claims provide:

- Claim
- Assessment
- Evidence
- Explanation
- Source URL

Possible assessments:

Supported
Partially Supported
Unclear

Never mark a claim as verified without sufficient evidence.
""",

    "Visualization Agent": """
You are the Visualization Agent.

Analyze supplied data and recommend an appropriate visualization.

Never invent numerical data.

When numerical data is supplied, generate an appropriate chart.

When the user explicitly requests an image or picture,
use the image-generation system.
""",

    "Summary Agent": """
You are the Summary Agent.

Create a concise but complete summary.

Use:

Overview
Key Points
Important Findings
Conclusion

Preserve important facts.

Do not invent information.
""",

    "Orchestrator Agent": """
You are the Orchestrator Agent.

Coordinate specialist agents.

Determine which specialists are required.

Combine their outputs into one useful final answer.

Do not invent information.

Briefly explain the workflow.
"""
}


# ============================================================
# LLM
# ============================================================

def get_instruction(agent):

    return AGENT_INSTRUCTIONS.get(
        agent,
        "You are a professional AI assistant inside ResearchMind AI."
    )


def llm_call(
    system_prompt,
    user_prompt,
    max_tokens=2500
):

    response = groq_client.chat.completions.create(

        model="openai/gpt-oss-120b",

        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],

        temperature=0.3,
        max_tokens=max_tokens
    )

    return response.choices[0].message.content


# ============================================================
# WEB SEARCH
# ============================================================

def web_search(query, max_results=5):

    if not tavily_client:

        return (
            "Web search is unavailable because "
            "TAVILY_API_KEY is missing."
        )

    try:

        print("TAVILY: Searching:", query)

        result = tavily_client.search(
            query=query,
            search_depth="basic",
            max_results=max_results
        )

        results = result.get("results", [])

        if not results:
            return "No web results were found."

        formatted = []

        for item in results:

            formatted.append(
                f"Title: {item.get('title', '')}\n"
                f"URL: {item.get('url', '')}\n"
                f"Content: {item.get('content', '')[:1200]}"
            )

        return "\n\n".join(formatted)

    except Exception as exc:

        print("Tavily error:", repr(exc))

        return "Web search was temporarily unavailable."


# ============================================================
# KNOWLEDGE SEARCH
# ============================================================

def knowledge_search(query):

    if not knowledge_store:

        return "The private knowledge base is empty."

    words = set(
        word.lower()
        for word in re.findall(r"\w+", query)
        if len(word) > 2
    )

    matches = []

    for item in knowledge_store:

        combined_text = (
            f"{item['title']} {item['content']}"
        ).lower()

        score = sum(
            1
            for word in words
            if word in combined_text
        )

        if score > 0:

            matches.append(
                (score, item)
            )

    matches.sort(
        key=lambda x: x[0],
        reverse=True
    )

    if not matches:

        return (
            "No matching information was found "
            "in the private knowledge base."
        )

    output = []

    for _, item in matches[:5]:

        output.append(
            f"Title: {item['title']}\n"
            f"Content: {item['content']}"
        )

    return "\n\n".join(output)


# ============================================================
# PDF UPLOAD
# ============================================================

@app.post("/upload-pdf")
async def upload_pdf(
    file: UploadFile = File(...)
):

    if (
        not file.filename
        or not file.filename.lower().endswith(".pdf")
    ):

        raise HTTPException(
            status_code=400,
            detail="Only PDF files are allowed."
        )

    data = await file.read()

    if not data:

        raise HTTPException(
            status_code=400,
            detail="The PDF file is empty."
        )

    # 25 MB upload limit
    if len(data) > 25 * 1024 * 1024:

        raise HTTPException(
            status_code=413,
            detail="PDF is too large. Maximum size is 25 MB."
        )

    temp_path = (
        Path(tempfile.gettempdir())
        / f"researchmind_{uuid.uuid4()}.pdf"
    )

    temp_path.write_bytes(data)

    try:

        reader = PdfReader(str(temp_path))

        pages = [
            page.extract_text() or ""
            for page in reader.pages
        ]

        text = "\n\n".join(pages).strip()

        # ====================================================
        # OCR FALLBACK
        # ====================================================

        if not text and OCR_AVAILABLE:

            tesseract_path = Path(
                r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            )

            if tesseract_path.exists():

                pytesseract.pytesseract.tesseract_cmd = (
                    str(tesseract_path)
                )

            pdf = fitz.open(str(temp_path))

            ocr_pages = []

            try:

                for page_number, page in enumerate(pdf):

                    print(
                        f"OCR processing page "
                        f"{page_number + 1}/{len(pdf)}"
                    )

                    pix = page.get_pixmap(
                        matrix=fitz.Matrix(2, 2),
                        alpha=False
                    )

                    image = Image.frombytes(
                        "RGB",
                        [
                            pix.width,
                            pix.height
                        ],
                        pix.samples
                    )

                    ocr_text = pytesseract.image_to_string(
                        image,
                        lang="eng"
                    )

                    ocr_pages.append(ocr_text)

            finally:

                pdf.close()

            text = "\n\n".join(
                ocr_pages
            ).strip()

        if not text:

            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract readable text. "
                    "If this is a scanned PDF, install "
                    "Tesseract OCR."
                )
            )

        # ====================================================
        # STORE DOCUMENT
        # ====================================================

        document_id = str(
            uuid.uuid4()
        )

        pdf_store[document_id] = {

            "filename": file.filename,

            "text": text,

            "pages": len(reader.pages),

            "characters": len(text)
        }

        return {

            "document_id": document_id,

            "filename": file.filename,

            "pages": len(reader.pages),

            "characters": len(text),

            "message": (
                "PDF uploaded and extracted successfully."
            )
        }

    finally:

        try:

            temp_path.unlink()

        except Exception:

            pass


# ============================================================
# DOCUMENT STATUS
# ============================================================

@app.get("/documents/{document_id}")
def get_document(document_id: str):

    document = pdf_store.get(
        document_id
    )

    if not document:

        raise HTTPException(
            status_code=404,
            detail="Document not found."
        )

    return {

        "document_id": document_id,

        "filename": document["filename"],

        "pages": document["pages"],

        "characters": document["characters"],

        "available": True
    }


# ============================================================
# KNOWLEDGE BASE
# ============================================================

@app.post("/knowledge/add")
def add_knowledge(
    request: KnowledgeRequest
):

    title = request.title.strip()
    content = request.content.strip()

    if not title:

        raise HTTPException(
            status_code=400,
            detail="Knowledge title is required."
        )

    if not content:

        raise HTTPException(
            status_code=400,
            detail="Knowledge content is required."
        )

    item = {

        "id": str(uuid.uuid4()),

        "title": title,

        "content": content
    }

    knowledge_store.append(item)

    return {

        "message": "Knowledge added successfully.",

        "item": item,

        "items": len(knowledge_store)
    }


@app.get("/knowledge")
def get_knowledge():

    return {

        "items": knowledge_store,

        "count": len(knowledge_store)
    }


@app.delete("/knowledge/{knowledge_id}")
def delete_knowledge(
    knowledge_id: str
):

    for index, item in enumerate(
        knowledge_store
    ):

        if item["id"] == knowledge_id:

            knowledge_store.pop(index)

            return {

                "message": "Knowledge deleted successfully."
            }

    raise HTTPException(
        status_code=404,
        detail="Knowledge item not found."
    )


# ============================================================
# REPORT EXPORT
# ============================================================

@app.post("/export-report")
def export_report(
    request: ReportExportRequest
):

    export_format = (
        request.format
        .lower()
        .strip()
    )

    if export_format not in {
        "pdf",
        "docx",
        "md",
        "txt",
        "html"
    }:

        raise HTTPException(
            status_code=400,
            detail=(
                "Supported formats: "
                "pdf, docx, md, txt, html"
            )
        )

    safe_name = "".join(

        char
        if char.isalnum() or char in " _-"
        else "_"

        for char in request.title
    ).strip()

    if not safe_name:

        safe_name = "research_report"

    output_dir = (
        Path(tempfile.gettempdir())
        / "researchmind_exports"
    )

    output_dir.mkdir(
        exist_ok=True
    )

    # ========================================================
    # MARKDOWN
    # ========================================================

    if export_format == "md":

        path = (
            output_dir
            / f"{safe_name}.md"
        )

        path.write_text(
            request.content,
            encoding="utf-8"
        )

    # ========================================================
    # TXT
    # ========================================================

    elif export_format == "txt":

        path = (
            output_dir
            / f"{safe_name}.txt"
        )

        path.write_text(
            request.content,
            encoding="utf-8"
        )

    # ========================================================
    # HTML
    # ========================================================

    elif export_format == "html":

        path = (
            output_dir
            / f"{safe_name}.html"
        )

        html_content = (
            request.content
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>{request.title}</title>

<style>

body {{
    font-family: Arial, sans-serif;
    max-width: 900px;
    margin: 40px auto;
    padding: 20px;
    line-height: 1.7;
}}

pre {{
    white-space: pre-wrap;
}}

</style>

</head>

<body>

<h1>{request.title}</h1>

<pre>{html_content}</pre>

</body>
</html>
"""

        path.write_text(
            html,
            encoding="utf-8"
        )

    # ========================================================
    # DOCX
    # ========================================================

    elif export_format == "docx":

        path = (
            output_dir
            / f"{safe_name}.docx"
        )

        document = Document()

        document.add_heading(
            request.title,
            level=0
        )

        for line in request.content.splitlines():

            if line.startswith("# "):

                document.add_heading(
                    line[2:],
                    level=1
                )

            elif line.startswith("## "):

                document.add_heading(
                    line[3:],
                    level=2
                )

            elif line.startswith("### "):

                document.add_heading(
                    line[4:],
                    level=3
                )

            elif line.strip():

                document.add_paragraph(
                    line
                )

        document.save(path)

    # ========================================================
    # PDF
    # ========================================================

    else:

        path = (
            output_dir
            / f"{safe_name}.pdf"
        )

        styles = getSampleStyleSheet()

        document = SimpleDocTemplate(

            str(path),

            pagesize=A4,

            rightMargin=45,

            leftMargin=45,

            topMargin=45,

            bottomMargin=45
        )

        story = [

            Paragraph(
                request.title,
                styles["Title"]
            ),

            Spacer(1, 18)
        ]

        for line in request.content.splitlines():

            clean_line = line.strip()

            if not clean_line:

                story.append(
                    Spacer(1, 8)
                )

                continue

            safe_text = (
                clean_line
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

            if clean_line.startswith("# "):

                story.append(
                    Paragraph(
                        safe_text[2:],
                        styles["Heading1"]
                    )
                )

            elif clean_line.startswith("## "):

                story.append(
                    Paragraph(
                        safe_text[3:],
                        styles["Heading2"]
                    )
                )

            elif clean_line.startswith("### "):

                story.append(
                    Paragraph(
                        safe_text[4:],
                        styles["Heading3"]
                    )
                )

            else:

                story.append(
                    Paragraph(
                        safe_text,
                        styles["BodyText"]
                    )
                )

            story.append(
                Spacer(1, 6)
            )

        document.build(story)

    return FileResponse(

        path=str(path),

        filename=path.name,

        media_type="application/octet-stream"
    )


# ============================================================
# RESEARCH AGENT
# ============================================================

def research_agent(prompt):

    web = web_search(
        prompt,
        max_results=5
    )

    return llm_call(

        get_instruction(
            "Research Agent"
        ),

        f"""
USER REQUEST:

{prompt}

WEB RESEARCH:

{web}

Produce evidence-based research.

Include the supplied source URLs.

Do not invent sources.
""",

        max_tokens=1800
    )


# ============================================================
# PDF ANALYSIS AGENT
# ============================================================

def pdf_agent(
    prompt,
    document_id
):

    if not document_id:

        return (
            "📄 Please upload a PDF first. "
            "The frontend must send the returned "
            "document_id with the request."
        )

    document = pdf_store.get(
        document_id
    )

    if not document:

        return (
            "The requested PDF was not found. "
            "Please upload it again."
        )

    text = document["text"]

    # Protect the model context.
    text_for_ai = text[:50000]

    return llm_call(

        get_instruction(
            "PDF Analysis Agent"
        ),

        f"""
PDF FILE:

{document['filename']}

USER QUESTION:

{prompt}

EXTRACTED PDF CONTENT:

{text_for_ai}

Answer the user's question using only
the supplied PDF content.

If the answer is not available in the PDF,
say that clearly.
""",

        max_tokens=2200
    )


# ============================================================
# KNOWLEDGE AGENT
# ============================================================

def knowledge_agent(prompt):

    knowledge = knowledge_search(
        prompt
    )

    return llm_call(

        get_instruction(
            "Knowledge Agent"
        ),

        f"""
USER REQUEST:

{prompt}

PRIVATE KNOWLEDGE BASE:

{knowledge}

Use relevant private knowledge when available.

If the private knowledge base does not contain
the requested information, use general knowledge,
but clearly distinguish it when appropriate.

Answer naturally.
""",

        max_tokens=1600
    )


# ============================================================
# REPORT WRITER AGENT
# ============================================================

def report_writer_agent(prompt):

    web = web_search(
        prompt,
        max_results=5
    )

    return llm_call(

        get_instruction(
            "Report Writer Agent"
        ),

        f"""
RESEARCH TOPIC:

{prompt}

RESEARCH SOURCES:

{web}

Create a professional structured report.

Use Markdown.

Do not invent facts or references.
""",

        max_tokens=3000
    )


# ============================================================
# CITATION AGENT
# ============================================================

def citation_agent(prompt):

    web = web_search(
        prompt,
        max_results=5
    )

    return llm_call(

        get_instruction(
            "Citation Agent"
        ),

        f"""
USER REQUEST:

{prompt}

AVAILABLE SOURCES:

{web}

Create accurate citations and references
only from the supplied sources.

Do not invent references.
""",

        max_tokens=1800
    )


# ============================================================
# FACT CHECKER
# ============================================================

def fact_checker_agent(prompt):

    web = web_search(
        prompt,
        max_results=5
    )

    return llm_call(

        get_instruction(
            "Fact Checker Agent"
        ),

        f"""
CLAIM / QUESTION:

{prompt}

EVIDENCE FROM WEB:

{web}

Fact-check the claim.

Provide:

Claim
Assessment
Evidence
Explanation
Source URL

Do not claim verification without evidence.
""",

        max_tokens=2200
    )


# ============================================================
# VISUALIZATION HELPERS
# ============================================================

def _extract_json_object(text):

    text = text.strip()

    text = re.sub(
        r"^```(?:json)?\s*",
        "",
        text,
        flags=re.I
    )

    text = re.sub(
        r"\s*```$",
        "",
        text
    )

    match = re.search(
        r"\{.*\}",
        text,
        re.S
    )

    if not match:

        raise ValueError(
            "Visualization model did not return valid JSON."
        )

    return json.loads(
        match.group(0)
    )


# ============================================================
# HUGGING FACE IMAGE GENERATION
# ============================================================

def generate_image(prompt):

    if not HUGGINGFACE_API_KEY:

        raise RuntimeError(
            "HUGGINGFACE_API_KEY is missing "
            "from the .env file."
        )

    print(
        "IMAGE GENERATION: Using Hugging Face..."
    )

    print(
        "IMAGE PROMPT:",
        prompt
    )

    try:

        from huggingface_hub import (
            InferenceClient
        )

    except ImportError:

        raise RuntimeError(
            "huggingface_hub is not installed.\n"
            "Run:\n"
            "pip install -U huggingface_hub"
        )

    try:

        client = InferenceClient(

            provider="auto",

            api_key=HUGGINGFACE_API_KEY
        )

        image = client.text_to_image(

            prompt=prompt,

            model=(
                "black-forest-labs/"
                "FLUX.1-schnell"
            )
        )

    except Exception as exc:

        print(
            "Hugging Face image error:",
            repr(exc)
        )

        raise RuntimeError(
            "Hugging Face image generation failed: "
            f"{exc}"
        ) from exc

    output_dir = (

        Path(tempfile.gettempdir())
        / "researchmind_visualizations"
    )

    output_dir.mkdir(
        exist_ok=True
    )

    image_id = str(
        uuid.uuid4()
    )

    path = (
        output_dir
        / f"{image_id}.png"
    )

    image.save(
        path,
        format="PNG"
    )

    print(
        "IMAGE GENERATED:",
        image_id
    )

    return {

        "text": (
            "## 🖼️ Image Generated\n\n"
            "Your image has been generated successfully."
        ),

        "visualization_id":
            image_id,

        "visualization_url":
            f"/visualizations/{image_id}"
    }


# ============================================================
# VISUALIZATION AGENT
# ============================================================

def extract_latest_user_request(prompt):
    """
    Extract only the latest user request from the conversation
    sent by the frontend.

    This prevents previous image/chart requests from affecting
    the current Visualization Agent request.
    """

    if not prompt:
        return ""

    # Look for the latest explicit user message.
    matches = re.findall(
        r"(?:LATEST USER MESSAGE|USER)\s*:\s*(.*?)(?=\n(?:USER|ASSISTANT|LATEST USER MESSAGE)\s*:|$)",
        prompt,
        flags=re.IGNORECASE | re.DOTALL
    )

    if matches:
        return matches[-1].strip()

    # Fallback: use the entire prompt.
    return prompt.strip()


def visualization_agent(prompt):

    # ========================================================
    # GET ONLY CURRENT REQUEST
    # ========================================================

    current_request = extract_latest_user_request(
        prompt
    )

    lower_prompt = current_request.lower().strip()

    print("\n========================================")
    print("VISUALIZATION AGENT")
    print("CURRENT REQUEST:")
    print(current_request)
    print("========================================")

    # ========================================================
    # CHART DETECTION
    # ========================================================

    chart_keywords = [
        "bar chart",
        "bar graph",
        "bar plot",
        "line chart",
        "line graph",
        "line plot",
        "pie chart",
        "pie graph",
        "pie plot",
        "scatter chart",
        "scatter graph",
        "scatter plot",
        "histogram",
        "chart",
        "graph",
        "plot",
    ]

    wants_chart = any(
        keyword in lower_prompt
        for keyword in chart_keywords
    )

    # ========================================================
    # IMAGE DETECTION
    # ========================================================

    image_keywords = [
        "generate an image",
        "generate a image",
        "generate image",
        "create an image",
        "create a image",
        "create image",
        "make an image",
        "make a image",
        "make image",
        "generate a picture",
        "generate picture",
        "create a picture",
        "create picture",
        "make a picture",
        "make picture",
        "draw an image",
        "draw a image",
        "draw an",
        "draw a",
        "show me an image",
        "show me a image",
        "show me an picture",
        "show me a picture",
        "image of",
        "picture of",
        "photo of",
    ]

    wants_image = any(
        keyword in lower_prompt
        for keyword in image_keywords
    )

    # ========================================================
    # IMPORTANT:
    # CHART HAS PRIORITY OVER IMAGE
    # ========================================================

    if wants_chart:

        print(
            "VISUALIZATION MODE: CHART"
        )

        # ----------------------------------------------------
        # Ask LLM for chart structure
        # ----------------------------------------------------

        spec_prompt = f"""
You are a data visualization parser.

The user wants a CHART.

Extract ONLY numerical/category data explicitly present
in the current user request.

NEVER invent values.

CURRENT USER REQUEST:

{current_request}

Return ONLY valid JSON.

For a bar chart, line chart, or pie chart:

{{
    "chart_type": "bar",
    "title": "Chart title",
    "x_label": "Category",
    "y_label": "Value",
    "labels": ["A", "B", "C"],
    "values": [10, 20, 30]
}}

For a scatter plot:

{{
    "chart_type": "scatter",
    "title": "Chart title",
    "x_label": "X",
    "y_label": "Y",
    "x_values": [1, 2, 3],
    "y_values": [10, 20, 30]
}}

Allowed chart_type values:

bar
line
pie
scatter

IMPORTANT:

- Use ONLY data from the current request.
- Do not use previous conversation messages.
- Do not invent missing numbers.
- Return JSON only.
"""

        try:

            raw = llm_call(
                get_instruction(
                    "Visualization Agent"
                ),
                spec_prompt,
                max_tokens=1200
            )

            print(
                "CHART SPEC RAW:",
                raw
            )

            spec = _extract_json_object(
                raw
            )

        except Exception as exc:

            print(
                "Chart specification error:",
                repr(exc)
            )

            return {
                "text": (
                    "❌ I could not understand "
                    "the chart data.\n\n"
                    "Please provide the chart type "
                    "and numerical values clearly."
                ),
                "visualization_url": None
            }

        chart_type = str(
            spec.get(
                "chart_type",
                "bar"
            )
        ).lower().strip()

        title = str(
            spec.get(
                "title",
                "ResearchMind AI Visualization"
            )
        )

        output_dir = (
            Path(tempfile.gettempdir())
            / "researchmind_visualizations"
        )

        output_dir.mkdir(
            exist_ok=True
        )

        image_id = str(
            uuid.uuid4()
        )

        path = (
            output_dir
            / f"{image_id}.png"
        )

        # ----------------------------------------------------
        # CREATE CHART
        # ----------------------------------------------------

        plt.figure(
            figsize=(10, 6)
        )

        try:

            if chart_type == "bar":

                labels = spec.get(
                    "labels",
                    []
                )

                values = spec.get(
                    "values",
                    []
                )

                if (
                    not labels
                    or not values
                    or len(labels) != len(values)
                ):
                    raise ValueError(
                        "Bar chart requires matching labels and values."
                    )

                plt.bar(
                    labels,
                    values
                )

                plt.xlabel(
                    spec.get(
                        "x_label",
                        "Category"
                    )
                )

                plt.ylabel(
                    spec.get(
                        "y_label",
                        "Value"
                    )
                )

            elif chart_type == "line":

                labels = spec.get(
                    "labels",
                    []
                )

                values = spec.get(
                    "values",
                    []
                )

                if (
                    not labels
                    or not values
                    or len(labels) != len(values)
                ):
                    raise ValueError(
                        "Line chart requires matching labels and values."
                    )

                plt.plot(
                    labels,
                    values,
                    marker="o"
                )

                plt.xlabel(
                    spec.get(
                        "x_label",
                        "Category"
                    )
                )

                plt.ylabel(
                    spec.get(
                        "y_label",
                        "Value"
                    )
                )

            elif chart_type == "pie":

                labels = spec.get(
                    "labels",
                    []
                )

                values = spec.get(
                    "values",
                    []
                )

                if (
                    not labels
                    or not values
                    or len(labels) != len(values)
                ):
                    raise ValueError(
                        "Pie chart requires matching labels and values."
                    )

                plt.pie(
                    values,
                    labels=labels,
                    autopct="%1.1f%%"
                )

            elif chart_type == "scatter":

                x_values = spec.get(
                    "x_values",
                    []
                )

                y_values = spec.get(
                    "y_values",
                    []
                )

                if (
                    not x_values
                    or not y_values
                    or len(x_values) != len(y_values)
                ):
                    raise ValueError(
                        "Scatter plot requires matching X and Y values."
                    )

                plt.scatter(
                    x_values,
                    y_values
                )

                plt.xlabel(
                    spec.get(
                        "x_label",
                        "X"
                    )
                )

                plt.ylabel(
                    spec.get(
                        "y_label",
                        "Y"
                    )
                )

            else:

                raise ValueError(
                    f"Unsupported chart type: {chart_type}"
                )

            plt.title(
                title
            )

            plt.tight_layout()

            plt.savefig(
                path,
                dpi=160,
                bbox_inches="tight"
            )

        finally:

            plt.close()

        print(
            "CHART GENERATED:",
            path
        )

        return {

            "text": (
                "## 📊 Visualization Generated\n\n"
                f"**Chart:** {title}\n\n"
                f"**Type:** {chart_type}\n\n"
                "The chart was generated from "
                "the data provided in your current request."
            ),

            "visualization_id":
                image_id,

            "visualization_url":
                f"/visualizations/{image_id}"
        }

    # ========================================================
    # IMAGE GENERATION
    # ========================================================

    if wants_image:

        print(
            "VISUALIZATION MODE: IMAGE"
        )

        # IMPORTANT:
        # Only send the current request to Hugging Face.
        image_prompt = current_request

        try:

            return generate_image(
                image_prompt
            )

        except Exception as exc:

            print(
                "Image generation error:",
                repr(exc)
            )

            return {

                "text":
                    "❌ Image generation failed.\n\n"
                    f"{exc}",

                "visualization_url":
                    None
            }

    # ========================================================
    # NO VISUALIZATION TYPE DETECTED
    # ========================================================

    return {

        "text": (
            "🎨 **Visualization Agent**\n\n"
            "I can create either:\n\n"
            "🖼️ **Images** — for example:\n"
            "`Generate an image of a futuristic AI robot.`\n\n"
            "📊 **Charts** — for example:\n"
            "`Create a pie chart: Python 40%, Java 30%, C++ 30%.`"
        ),

        "visualization_url":
            None
    }

# ============================================================
# SUMMARY AGENT
# ============================================================

def summary_agent(prompt):

    return llm_call(

        get_instruction(
            "Summary Agent"
        ),

        f"""
CONTENT TO SUMMARIZE:

{prompt}

Create a concise but complete summary.

Use:

Overview
Key Points
Important Findings
Conclusion
""",

        max_tokens=1600
    )


# ============================================================
# ORCHESTRATOR
# ============================================================

def orchestrator_agent(prompt):

    print("ORCHESTRATOR: starting compact workflow")

    # Keep the original request small so chained Groq calls
    # stay below the free/on-demand TPM limit.
    user_request = (prompt or "").strip()[:3500]

    # ========================================================
    # 1. RESEARCH
    # ========================================================

    research = research_agent(
        user_request
    )

    # Keep only the most useful portion for the next agents.
    research_short = research[:5000]

    # ========================================================
    # 2. FACT CHECK
    # ========================================================

    fact_check = fact_checker_agent(
        f"""
USER REQUEST:
{user_request}

RESEARCH EVIDENCE:
{research_short}

Fact-check only the important claims.
Keep the response concise.
"""
    )

    fact_check_short = fact_check[:3000]

    # ========================================================
    # 3. CITATIONS
    # ========================================================

    citations = citation_agent(
        f"""
USER REQUEST:
{user_request}

RESEARCH SOURCES:
{research_short}

Create concise citations using ONLY the supplied sources.
"""
    )

    citations_short = citations[:2500]

    # ========================================================
    # 4. FINAL REPORT
    # ========================================================

    final_prompt = f"""
Create a concise final research answer.

USER REQUEST:
{user_request}

RESEARCH:
{research_short}

FACT CHECK:
{fact_check_short}

CITATIONS:
{citations_short}

Requirements:
- Use only the supplied evidence.
- Do not invent facts or sources.
- Clearly mention uncertainty.
- Use Markdown headings and bullet points.
- Keep the final answer concise.
"""

    report = llm_call(
        get_instruction(
            "Report Writer Agent"
        ),
        final_prompt,
        max_tokens=1800
    )

    return f"""
## ⚡ Orchestrator Workflow

The request was processed through:

1. 🔎 Research Agent
2. ✅ Fact Checker Agent
3. 🔗 Citation Agent
4. 📝 Report Writer Agent

## Final Result

{report}

## Fact-Checking Notes

{fact_check_short}

## References

{citations_short}
"""


# ============================================================
# MAIN AGENT ROUTER
# ============================================================

@app.post("/run-agent")
def run_agent_endpoint(
    request: AgentRequest
):

    print(
        "\n========================================"
    )

    print(
        "Agent:",
        request.agent
    )

    print(
        "Question:",
        request.prompt[:300]
    )

    print(
        "Document:",
        request.document_id
    )

    print(
        "========================================"
    )

    try:

        if request.agent == "Research Agent":

            result = {
                "text":
                    research_agent(
                        request.prompt
                    )
            }

        elif request.agent == "PDF Analysis Agent":

            result = {
                "text":
                    pdf_agent(
                        request.prompt,
                        request.document_id
                    )
            }

        elif request.agent == "Knowledge Agent":

            result = {
                "text":
                    knowledge_agent(
                        request.prompt
                    )
            }

        elif request.agent == "Report Writer Agent":

            result = {
                "text":
                    report_writer_agent(
                        request.prompt
                    )
            }

        elif request.agent == "Citation Agent":

            result = {
                "text":
                    citation_agent(
                        request.prompt
                    )
            }

        elif request.agent == "Fact Checker Agent":

            result = {
                "text":
                    fact_checker_agent(
                        request.prompt
                    )
            }

        elif request.agent == "Visualization Agent":

            result = visualization_agent(
                request.prompt
            )

        elif request.agent == "Summary Agent":

            result = {
                "text":
                    summary_agent(
                        request.prompt
                    )
            }

        elif request.agent == "Orchestrator Agent":

            result = {
                "text":
                    orchestrator_agent(
                        request.prompt
                    )
            }

        else:

            result = {

                "text":
                    llm_call(
                        (
                            "You are a professional "
                            "ResearchMind AI assistant."
                        ),

                        request.prompt
                    )
            }

        return {

            "agent":
                request.agent,

            "response":
                result.get(
                    "text",
                    "No response received."
                ),

            "visualization_url":
                result.get(
                    "visualization_url"
                )
        }

    except Exception as exc:

        print(
            "AI processing error:",
            repr(exc)
        )

        return {

            "agent":
                request.agent,

            "response":
                (
                    "⚠️ AI processing error.\n\n"
                    f"{exc}"
                ),

            "visualization_url":
                None
        }


# ============================================================
# VISUALIZATION FILE
# ============================================================

@app.get(
    "/visualizations/{image_id}"
)
def get_visualization(
    image_id: str
):

    path = (

        Path(tempfile.gettempdir())

        / "researchmind_visualizations"

        / f"{image_id}.png"
    )

    if not path.exists():

        raise HTTPException(
            status_code=404,
            detail="Visualization not found."
        )

    return FileResponse(
        str(path),
        media_type="image/png",
        filename=path.name
    )


# ============================================================
# MULTI-AGENT
# ============================================================

@app.post("/run-multi-agent")
def run_multi_agent(
    request: MultiAgentRequest
):

    selected = list(
        dict.fromkeys(
            request.agents
        )
    )

    if not selected:

        raise HTTPException(
            status_code=400,
            detail="Select at least one agent."
        )

    if len(selected) > 6:

        raise HTTPException(
            status_code=400,
            detail="Select at most 6 agents at once."
        )

    outputs = []

    for agent_name in selected:

        if agent_name == "Orchestrator Agent":

            result = orchestrator_agent(
                request.prompt
            )

            outputs.append({

                "agent":
                    agent_name,

                "response":
                    result,

                "visualization_url":
                    None
            })

            continue

        elif agent_name == "Research Agent":

            result = research_agent(
                request.prompt
            )

        elif agent_name == "PDF Analysis Agent":

            result = pdf_agent(
                request.prompt,
                request.document_id
            )

        elif agent_name == "Knowledge Agent":

            result = knowledge_agent(
                request.prompt
            )

        elif agent_name == "Report Writer Agent":

            result = report_writer_agent(
                request.prompt
            )

        elif agent_name == "Citation Agent":

            result = citation_agent(
                request.prompt
            )

        elif agent_name == "Fact Checker Agent":

            result = fact_checker_agent(
                request.prompt
            )

        elif agent_name == "Visualization Agent":

            result = visualization_agent(
                request.prompt
            )

            if isinstance(
                result,
                dict
            ):

                outputs.append({

                    "agent":
                        agent_name,

                    "response":
                        result["text"],

                    "visualization_url":
                        result.get(
                            "visualization_url"
                        )
                })

                continue

        elif agent_name == "Summary Agent":

            result = summary_agent(
                request.prompt
            )

        else:

            continue

        outputs.append({

            "agent":
                agent_name,

            "response":
                result,

            "visualization_url":
                None
        })

    combined = "\n\n".join(

        f"### {item['agent']}\n"
        f"{item['response']}"

        for item in outputs
    )

    final = llm_call(

        get_instruction(
            "Orchestrator Agent"
        ),

        f"""
Coordinate the selected specialist outputs
into one final answer.

USER REQUEST:

{request.prompt}

SELECTED AGENTS:

{', '.join(selected)}

SPECIALIST OUTPUTS:

{combined}

Do not invent information.

Preserve useful specialist findings.

Clearly mention uncertainty.
""",

        max_tokens=4500
    )

    visualizations = [

        item["visualization_url"]

        for item in outputs

        if item.get(
            "visualization_url"
        )
    ]

    return {

        "agents":
            selected,

        "responses":
            outputs,

        "response":
            final,

        "visualization_urls":
            visualizations
    }


# ============================================================
# HEALTH CHECK
# ============================================================

@app.get("/")
def root():

    return {

        "status":
            "online",

        "application":
            "ResearchMind AI",

        "message":
            "ResearchMind AI backend is running",

        "agents":
            9,

        "pdf_upload":
            True,

        "pdf_analysis":
            True,

        "report_export":
            True,

        "knowledge_base":
            True,

        "research":
            bool(TAVILY_API_KEY),

        "orchestrator":
            True,

        "multi_agent":
            True,

        "visualization_images":
            True,

        "huggingface_image_generation":
            bool(HUGGINGFACE_API_KEY),

        "ocr":
            OCR_AVAILABLE
    }


# ============================================================
# RUN SERVER
# ============================================================

if __name__ == "__main__":

    import uvicorn

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        reload=False
    )